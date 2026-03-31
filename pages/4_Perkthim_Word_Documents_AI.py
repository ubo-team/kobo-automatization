import streamlit as st
from io import BytesIO
import os
import re
import time
import google.generativeai as genai
from docx import Document

st.set_page_config(page_title="Përkthe Word Dokumente me AI", layout="centered")


logo_svg_path = "UBO-Logo.svg"
with st.sidebar:
    if os.path.exists(logo_svg_path):
        with open(logo_svg_path, "r", encoding="utf-8") as f:
            svg_logo = f.read()
        st.markdown(
            f'<div style="display:flex;justify-content:center;margin:15px 0;"><div style="width:150px;">{svg_logo}</div></div>',
            unsafe_allow_html=True
        )
    st.markdown("""
        <style>
        [data-testid="stSidebar"] img {
            display: block;
            margin-left: auto;
            margin-right: auto;
            margin-top: 15px;
            margin-bottom: 15px;
        }
        </style>
    """, unsafe_allow_html=True)


GEMINI_TRANSLATION_API_KEY = st.secrets["GEMINI_TRANSLATION_API_KEY"]
genai.configure(api_key=GEMINI_TRANSLATION_API_KEY)

GEMINI_PRICING = {
    "models/gemini-2.5-flash": {
        "inputPer1MTokens":  0.30,
        "outputPer1MTokens": 2.50,
    },
}


def calculate_gemini_cost(prompt_tokens: int, completion_tokens: int, model: str) -> float:
    pricing = GEMINI_PRICING.get(model, GEMINI_PRICING["models/gemini-2.5-flash"])
    input_rate  = pricing["inputPer1MTokens"]
    output_rate = pricing["outputPer1MTokens"]
    input_cost  = ((prompt_tokens     or 0) / 1_000_000) * input_rate
    output_cost = ((completion_tokens or 0) / 1_000_000) * output_rate
    return input_cost + output_cost


LANGUAGE_OPTIONS_UI = {
    "Gjuha Shqipe": "sq",
    "Gjuha Angleze": "en",
    "Gjuha Serbe": "sr",
    "Gjuha Maqedonase": "mk",
    "Gjuha Boshnjake": "bs"
}

LANG_NAMES = {
    "sq": "Albanian",
    "en": "English",
    "sr": "Serbian (Latin script)",
    "mk": "Macedonian (Latin script)",
    "bs": "Bosnian"
}

MODEL_NAME = "gemini-3.1-flash-lite-preview"
gemini_model = genai.GenerativeModel(MODEL_NAME)


BATCH_SIZE = 50


def translate_batch(texts, from_lang, to_lang):
    """Translate a batch of texts in one API call. Returns (translations_dict, in_tok, out_tok, error)."""
    from_name = LANG_NAMES.get(from_lang, from_lang)
    to_name = LANG_NAMES.get(to_lang, to_lang)

    numbered_texts = "\n".join(f"[{j+1}] {t}" for j, t in enumerate(texts))
    prompt = (
        f"{from_name} to {to_name}. Reply [N] translation only.\n\n{numbered_texts}"
    )

    response = gemini_model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(temperature=0.1, max_output_tokens=4096),
    )
    in_tok = getattr(response.usage_metadata, "prompt_token_count", 0) or 0
    out_tok = getattr(response.usage_metadata, "candidates_token_count", 0) or 0

    translations = {}
    for line in response.text.strip().split("\n"):
        m = re.match(r"\[(\d+)\]\s*(.*)", line.strip())
        if m:
            translations[int(m.group(1))] = m.group(2).strip()

    return translations, in_tok, out_tok


def translate_docx_in_place(doc, from_lang, to_lang):
    total_in, total_out = 0, 0
    errors = []

    # Collect all translatable runs with their references
    run_entries = []  # (run_reference, original_text)

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if text:
                run_entries.append((run, text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if text:
                            run_entries.append((run, text))

    if not run_entries:
        return doc, 0, 0, []

    progress = st.progress(0, text="Duke përkthyer... 0%")

    for batch_start in range(0, len(run_entries), BATCH_SIZE):
        batch = run_entries[batch_start:batch_start + BATCH_SIZE]
        texts = [text for _, text in batch]

        try:
            translations, in_tok, out_tok = translate_batch(texts, from_lang, to_lang)
            total_in += in_tok
            total_out += out_tok

            for j, (run, _) in enumerate(batch):
                if (j + 1) in translations:
                    run.text = translations[j + 1]
        except Exception as e:
            errors.append(str(e))

        done = min(batch_start + BATCH_SIZE, len(run_entries))
        pct = done / len(run_entries)
        progress.progress(pct, text=f"Duke përkthyer... {done}/{len(run_entries)}")

    progress.empty()
    return doc, total_in, total_out, errors


st.title("Fillo me Përkthimin e Pyetësorëve")

with st.expander("Testo API Key"):
    if st.button("Testo Gemini API"):
        try:
            test_response = gemini_model.generate_content("Translate 'Hello' to Albanian. Return ONLY the translation.")
            st.success(f"API funksionon! Pergjigja: {test_response.text.strip()}")
        except Exception as e:
            st.error(f"API nuk funksionon: {e}")

uploaded_file = st.file_uploader("Ngarko dokumentin (vetëm Word)", type=["docx"])

if uploaded_file:
    from_lang_label = st.selectbox("Gjuha Burimore", list(LANGUAGE_OPTIONS_UI.keys()), key="word_lang_from")
    to_lang_label = st.selectbox("Gjuha për Përkthim", list(LANGUAGE_OPTIONS_UI.keys()), key="word_lang_to")
    from_lang = LANGUAGE_OPTIONS_UI[from_lang_label]
    to_lang = LANGUAGE_OPTIONS_UI[to_lang_label]

    if st.button("Përkthe Word Dokumentin"):
        doc = Document(uploaded_file)
        translated_doc, total_in, total_out, errors = translate_docx_in_place(doc, from_lang, to_lang)

        output = BytesIO()
        translated_doc.save(output)
        output.seek(0)

        if errors:
            st.error(f"Ka pasur {len(errors)} gabime. Gabimi i parë: {errors[0]}")
        else:
            st.success("Përkthimi përfundoi me sukses!")

        model_id = f"models/{MODEL_NAME}"
        cost = calculate_gemini_cost(total_in, total_out, model_id)
        st.info(
            f"**Kostoja:**  \n"
            f"Input tokens: **{total_in:,}** | Output tokens: **{total_out:,}**  \n"
            f"Kostoja: **${cost:.4f}**"
        )

        st.download_button(
            label="Shkarko dokumentin e përkthyer (Word)",
            data=output,
            file_name=f"translated_{uploaded_file.name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
