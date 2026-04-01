import pandas as pd
import re
import streamlit as st
from docx import Document
from collections import defaultdict
from difflib import get_close_matches
import os

QUESTION_PATTERN = re.compile(
    r"\[ *(single|multiple|open|text|numeric|matrix(?: [a-z0-9]*)*|multiple matrix|scale|other) *\]",
    re.IGNORECASE,
)

HINT_PATTERN = re.compile(r"\[ *Hint: *(.*?) *\]", re.IGNORECASE) 

def clean_question_text(text):
    return QUESTION_PATTERN.sub("", text).strip()

def get_numbering(para):
    for run in para.runs:
        if run.text.strip():
            return run.text.strip()
    return ""

def is_list_option(para):
    return para.style and para.style.name.lower().startswith("list")

def extract_matrix_table(table, data, parent_qid, parent_qtext):
    rows = table.rows
    if not rows or len(rows) < 2:
        return
    headers = [cell.text.strip() for cell in rows[0].cells[1:] if cell.text.strip()]
    data.append({"Question ID": parent_qid, "Question Text": parent_qtext, "Option ID": None, "Option Text": None, "hint": ""})
    for row_index, row in enumerate(rows[1:], start=1):
        cells = row.cells
        subquestion_text = cells[0].text.strip()
        if not subquestion_text:
            continue
        sub_qid = f"{parent_qid}_{row_index}"
        data.append({"Question ID": sub_qid, "Question Text": subquestion_text, "Option ID": None, "Option Text": None, "hint": ""})
        for col_index, option_text in enumerate(headers, start=1):
            if option_text:
                option_qid = f"{sub_qid}.{col_index}"
                data.append({"Question ID": sub_qid, "Question Text": "", "Option ID": option_qid, "Option Text": option_text, "hint": ""})

def extract_from_docx_to_excel(docx_file):
    doc = Document(docx_file)
    data, question_counter, table_index = [], 0, 0
    paragraphs, para_index = list(doc.paragraphs), 0
    skip_options, last_matrix_qid, last_matrix_qtext = False, None, None

    while para_index < len(paragraphs):
        para = paragraphs[para_index]
        text = para.text.strip()

        if not text:
            para_index += 1
            continue

        if QUESTION_PATTERN.search(text) or re.match(r"^Q\d+\.", text):
            question_counter += 1
            qid = f"Q{question_counter}"

            hint_match = HINT_PATTERN.search(text)
            hint_text = hint_match.group(1).strip() if hint_match else ""

            cleaned = HINT_PATTERN.sub("", clean_question_text(text)).strip()

            data.append({
                "Question ID": qid,
                "Question Text": cleaned,
                "Option ID": None,
                "Option Text": None,
                "hint": hint_text
            })

            if "matrix" in text.lower():
                if table_index < len(doc.tables):
                    last_matrix_qid, last_matrix_qtext = qid, cleaned
                    extract_matrix_table(doc.tables[table_index], data, qid, cleaned)
                    table_index += 1
                skip_options = True
            elif "scale" in text.lower():
                skip_options = True
            else:
                skip_options = False

        elif is_list_option(para) and not skip_options and question_counter:
            prefix = get_numbering(para)
            full_option = f"{prefix} {text}" if prefix and not text.startswith(prefix) else text
            qid = f"Q{question_counter}"
            option_count = len([d for d in data if d['Question ID'] == qid and d['Option ID']])
            data.append({
                "Question ID": qid,
                "Question Text": None,
                "Option ID": f"{qid}_option_{option_count+1}",
                "Option Text": full_option,
                "hint": ""  # Options don't have hints
            })

        para_index += 1

    while table_index < len(doc.tables):
        table = doc.tables[table_index]
        first_col = [row.cells[0].text.strip() for row in table.rows[1:] if row.cells and row.cells[0].text.strip()]
        if any(re.match(r"\d+(\.\d+)?", cell) for cell in first_col):
            parent_qid = last_matrix_qid or f"Q{question_counter+1}"
            extract_matrix_table(table, data, parent_qid, last_matrix_qtext or "Matrix Question")
        table_index += 1

    return pd.DataFrame(data)

st.title("Përkthimi i dokumenteve zyrtare")
mode = st.radio("Zgjidh mënyrën:", ["Ngarko DOCX", "Ngarko XLSForm"])

if mode == "Ngarko DOCX":
    docx_file = st.file_uploader("Ngarko dokumentin Word (.docx)", type=["docx"])
    if docx_file:
        extracted_df = extract_from_docx_to_excel(docx_file)
        st.success("Dokumenti DOCX u ekstraktua me sukses!")
        st.dataframe(extracted_df.head())
        extracted_df.to_excel("cleaned_output.xlsx", index=False)
        st.download_button("Shkarko Excelin e ekstraktuar", data=open("cleaned_output.xlsx", "rb").read(), file_name="cleaned_output.xlsx")

elif mode == "Ngarko XLSForm":
    original_file = st.file_uploader("Ngarko XLSForm-in origjinal (Excelin)", type=["xlsx"])
    translated_file = st.file_uploader("Ngarko Excelin me përkthimin e pastruar", type=["xlsx"])

    if original_file and translated_file:
        import os
        from collections import defaultdict
        from difflib import get_close_matches

        xls = pd.ExcelFile(original_file)
        survey_df = xls.parse("survey")
        choices_df = xls.parse("choices")
        settings_df = xls.parse("settings")
        translated_df = pd.read_excel(translated_file)

        survey_df.columns = survey_df.columns.str.strip()
        choices_df.columns = choices_df.columns.str.strip()
        settings_df.columns = settings_df.columns.str.strip()

        label_columns = [col for col in survey_df.columns if col.startswith("label::")]
        if not label_columns:
            st.error("XLSForm nuk ka kolona 'label::'. Sigurohu që Exceli ka kolona si 'label::Albanian (1)', 'label::Serbian (2)', etj.")
            st.stop()
        from_label = st.selectbox("Zgjidh kolonën në Excel prej nga do të përkthehet:", label_columns).strip()
        to_label = st.selectbox("Zgjidh kolonën në Excel ku do të vendoset përkthimi:", label_columns).strip()

        hint_columns = [col for col in survey_df.columns if col.startswith("hint::")]
        if hint_columns:
            to_hint_col = st.selectbox("Zgjidh kolonën e hint-it ku do të vendoset (p.sh. Albanian ose Serbian):", hint_columns)
        else:
            to_hint_col = None

        LANG_OPTIONS = {
            "Gjuha Shqipe": "al",
            "Gjuha Angleze": "en",
            "Gjuha Serbe": "sr",
        }
        from_lang_label = st.selectbox("Gjuha burimore:", list(LANG_OPTIONS.keys()), key="from_lang_zyrt")
        to_lang_label = st.selectbox("Gjuha e përkthimit:", list(LANG_OPTIONS.keys()), key="to_lang_zyrt")
        from_lang = LANG_OPTIONS[from_lang_label]
        to_lang = LANG_OPTIONS[to_lang_label]

        def clean_label(val):
            if pd.isna(val): return ""
            text = str(val).strip().lower()
            text = re.sub(r"[\u2013\u2014-]", "-", text)
            text = re.sub(r"\s+", " ", text)
            return text

        def capitalize_first(text):
            return text[0].upper() + text[1:] if text else text

        def build_translation_dictionaries():
            manual_al_to_sr, manual_sr_to_al = {}, {}
            manual_al_to_en, manual_en_to_al = {}, {}
            manual_sr_to_en, manual_en_to_sr = {}, {}

            core_pairs = [
                ("GPS", "GPS", "GPS"),
                ("Anketuesi_ja", "Anketar/e", "Enumerator"),
                ("A pranoni të merrni pjesë në anketë?", "Da li se slažete da učestvujete u anketi?", "Do you agree to participate in the survey?"),
                ("Arsyet e refuzimit", "Razlozi za odbijanje", "Reasons for refusal"),
                ("Tjetër, specifiko", "Drugo, navedite", "Other, specify"),
                ("Po", "Da", "Yes"),
                ("Jo", "Ne", "No"),
                ("Tjetër. Çka?", "Drugo. Šta?", "Other. What?"),
                ("Tjetër, ju lutem specifikoni", "Drugo, navedite", "Other, please specify")
            ]

            likert_pairs = [
                ("1-Aspak i kënaqur", "1-Uopšte nisam zadovoljan/na", "1-Not at all satisfied"),
                ("5-Plotësisht i kënaqur", "5-Potpuno zadovoljan/zadovoljna", "5-Completely satisfied"),

                ("1-Aspak nuk pajtohem", "1-Uopšte se ne slažem", "1-Strongly disagree"),
                ("5-Plotësisht pajtohem", "5-Potpuno se slažem", "5-Strongly agree"),

                ("1– Aspak efektive", "1 – Uopšte efektivno", "1-Not effective at all"),
                ("5– Plotësisht efektive", "5 – Potpuno efektivno", "5-Completely effective"),

                ("1-Aspak e sigurtë", "1-Uopšte nije bezbedno", "1-Not safe at all"),
                ("5-Plotësisht e sigurtë", "5-Potpuno je bezbedno", "5-Completely safe"),

                ("1-Aspak meritore", "1-Nimalo zaslužne", "1-Not deserving at all"),
                ("5-Plotësisht meritore", "5-Potpuno zaslužne", "5-Completely deserving"),

                ("Shumë negative", "Veoma negativno", "Very negative"),
                ("Shumë pozitive", "Veoma pozitivno", "Very positive"),

                ("1 – aspak i mirë", "1 – uopšte nije dobar", "1-Not good at all"),
                ("5 – shumë i mirë", "5 – veoma dobar", "5-Very good"),

                ("88 – Refuzoj të përgjigjem", "Odbijam odgovoriti", "88-Refuse to answer"),
                ("Refuzoj të përgjigjem", "Odbijam odgovoriti", "Refuse to answer")
            ]
        
            demographic_pairs = [
                ("D1. (GJINIA)", "D1. (ROD/POL)", "D1. (GENDER)"),
                ("D2. (MOSHA) (vjet)", "D2. (STAROST) (godine)", "D2. (AGE) (years)"),
                ("D3. (STATUSI MARTESOR)  Aktualisht Ju jeni...", "D3. (BRAČNO STANJE) Trenutno vi ste…", "D3. (MARITAL STATUS) Currently you are..."),
                ("D4.  (EDUKIMI)  Sa vite shkollë i keni kryer?", "D4. (OBRAZOVANJE) Koliko godina škole ste završili?", "D4. (EDUCATION) How many years of schooling have you completed?"),
                ("D5.  (PËRKATËSIA ETNIKE)  Cili është nacionaliteti Juaj/cilit grup i takoni?", "D5. (ETNIČKA PRIPADNOST) Koja je vaša etnička pripadnost/kojoj grupi pripadate?", "D5. (ETHNICITY) What is your nationality/which group do you belong to?"),
                ("Tjetër. Cili?", "Drugo. Koja?", "Other. Which?"),
                ("D6. (FAMILJA)  Sa anëtarë i ka familja Juaj?", "D6. (PORODICA) Koliko članova ima vaša porodica?", "D6. (FAMILY) How many members are in your family?"),
                ("D8. (TË ARDHURAT PERSONALE) A mund të na tregoni se sa kanë qenë të ardhurat personale në muajin e fundit?", 
                "D8. (LIČNI PRIHODI) Da li nam možete reći koliki su bili vaši lični prihodi u zadnjem mesecu?", 
                "D8. (PERSONAL INCOME) Can you tell us what your personal income was last month?"),
                ("D9.  (TË ARDHURAT FAMILJARE) A mund të na tregoni se sa kanë qenë të ardhurat familjare në muajin e fundit?", 
                "D9. (PORODIČNI PRIHODI) Da li nam možete reći koliki je bio vaš porodični prihod u zadnjem mesecu?", 
                "D9. (HOUSEHOLD INCOME) Can you tell us what your household income was last month?"),
                ("D10. Komuna", "Opstina", "D10. Municipality"),
                ("D11.    VENDBANIMI", "D11. PREBIVALIŠTE", "D11. Residence"),
                ("Emri i lagjes", "Naziv komšiluka", "Neighborhood name"),
                ("Emri i fshatit", "Ime sela", "Village name"),
                ("Emri dhe mbiemri", "Ime i prezime", "Full name"),
                ("Numri i telefonit", "Broj telefona", "Phone number")
            ]

            extra_pairs = [("Mashkull", "Muško", "Male"), ("Femër", "Žensko", "Female")]
            reason_pairs = [
                ("Mungesa e kohës", "Nedostatak vremena", "Lack of time"),
                ("Jo i interesuar", "Nije zainteresovan", "Not interested"),
                ("Mbrojtja e të dhënave, përdorimi i të drejtës së privatësisë", 
                "Zaštita podataka, korišćenje politike privatnosti", 
                "Data protection, use of privacy rights"),
                ("Nuk beson në sondazhe", "Ne veruje u ankete", "Does not believe in surveys"),
                ("Të tjera (nuk di të përgjigjet, kushtet e motit, frikë nga pyetjet)", 
                "Ostalo (ne zna da odgovori, vremenski uslovi, strah od pitanja)", 
                "Other (don’t know how to answer, weather conditions, fear of questions)"),
                ("Problemet e shëndetit", "Zdravstveni problemi", "Health problems"),
                ("Moshë më e vjetër", "Starije godine", "Older age"),
                ("Nuk i pëlqen subjekti i kërkimit", "Ne voli temu istraživanja", "Does not like research topic"),
                ("Ka pasur një përvojë të keqe me sondazhet", 
                "Imao/la je loše iskustvo sa anketama", 
                "Had a bad experience with surveys"),
                ("Asnjë arsye", "Nema razloga", "No reason")
            ]

            demographic_pairs_answers = [
                ("Mashkull", "Muško", "Male"),
                ("Femër", "Žensko", "Female"),
                ("I/ e martuar", "Oženjen/Udata", "Married"),
                ("I/ e pamartuar", "Neoženjen/Neudata", "Single"),
                ("I/ e ndarë", "Razveden/a", "Divorced"),
                ("I/e vej", "Udovac/udovica", "Widowed"),
                ("Disa vite të shkollës fillore", "Nekoliko godina osnovne škole", "Some years of primary school"),
                ("Shkolla fillore", "Osnovna škola", "Primary school"),
                ("Disa vite të shkollës së mesme", "Nekoliko godina srednje škole", "Some years of secondary school"),
                ("Shkolla e mesme", "Srednja škola", "Secondary school"),
                ("Student", "Student", "Student"),
                ("Fakultet", "Fakultet", "University"),
                ("Magjistraturë ose Doktoraturë", "Magistratura ili", "Masters or Doctorate"),
                ("Shqiptar", "Albanska", "Albanian"),
                ("Serb", "Srpska", "Serbian"),
                ("Boshnjak", "Bosanska", "Bosniak"),
                ("Goran", "Goranska", "Gorani"),
                ("Turk", "Turska", "Turkish"),
                ("Rom", "Romska", "Roma"),
                ("Ashkali", "Aškalijska", "Ashkali"),
                ("Egjiptas", "Egipatska", "Egyptian"),
                ("Tjetër. Cili?", "Drugo. Koja?", "Other. Which?"),
                ("DK/PP", "Ne znam/Bez odgovora", "Don't know/No answer"),
            ]

            income_pairs = [
                ("Deri 150 euro", "Do 150 evra", "Up to 150 euros"),
                ("151-300 euro", "151-300 evra", "151-300 euros"),
                ("301-450 euro", "301-450 evra", "301-450 euros"),
                ("451-600 euro", "451-600 evra", "451-600 euros"),
                ("601-750 euro", "601-750 evra", "601-750 euros"),
                ("751-900 euro", "751-900 evra", "751-900 euros"),
                ("Mbi 900 euro", "Preko 900 evra", "Over 900 euros"),
                ("Nuk kam realizuar fare të ardhura", "Nisam ostvario/la nikakav prihod.", "I had no income"),
                ("Refuzon/PP", "Odbija/BO", "Refused/No answer")
            ]

            frequency_pairs = [
                ("Asnjëherë", "Nikad", "Never"),
                ("Rallë", "Retko", "Rarely"),
                ("Ndonjëherë", "Ponekad", "Sometimes"),
                ("Shpesh", "Često", "Often"),
                ("Gjithmonë", "Uvek", "Always")
            ]

            awareness_pairs = [
                ("Shumë i informuar", "Veoma informisani", "Very informed"),
                ("Deri diku i informuar", "Donekle informisani", "Somewhat informed"),
                ("Deri diku jo i informuar", "Donekle ne informisani", "Somewhat uninformed"),
                ("Aspak i informuar", "Potpuno ne informisani", "Not at all informed")
            ]

            satisfaction_pairs = [
                ("Shumë të kënaqur", "Veoma zadovoljni", "Very satisfied"),
                ("Deri diku i kënaqur", "Donekle zadovoljni", "Somewhat satisfied"),
                ("Deri diku jo i kënaqur", "Donekle nezadovoljni", "Somewhat dissatisfied"),
                ("Aspak i kënaqur", "Potpuno nezadovoljni", "Not at all satisfied"),
                ("Shumë i/e kënaqur", "Veoma zadovoljni", "Very satisfied"),
                ("I/e kënaqur", "Zadovoljni", "Satisfied"),
                ("I/e pakënaqur", "Nezadovoljni", "Dissatisfied"),
                ("Shumë i/e pakënaqur", "Veoma nezadovoljni", "Very dissatisfied"),
                ("Nuk e di/refuzoj të përgjigjem (mos e lexo)", 
                "Ne znam/Odbijam odgovoriti (nemojte čitati)", 
                "Don't know/Refuse to answer (do not read)")
            ]
            employment_pairs = [
                ("I papunësuar – duke kërkuar punë", "Nezaposlen/a – tražim posao", "Unemployed – seeking work"),
                ("I papunësuar – duke mos kërkuar punë", "Nezaposlen/a – ne tražim posao", "Unemployed – not seeking work"),
                ("I punësuar në sektorin publik", "Zaposlen/a u javnom sektoru", "Employed in public sector"),
                ("I punësuar në sektorin privat", "Zaposlen/a u privatnom sektoru", "Employed in private sector"),
                ("I punësuar kohë pas kohe", "Zaposlen/a s vremena na vreme", "Employed occasionally"),
                ("Pensionist", "Penzioner", "Pensioner"),
                ("Amvise", "Domaćica", "Housewife"),
                ("Student/ nxënës", "Student/učenik", "Student/Pupil"),
                ("Tjetër. Çka?", "Drugo. Šta?", "Other. What?")
            ]

            voting_pairs = [
                ("Gjithsesi do të votoja", "Svakako bih glasao/ala", "Definitely would vote"),
                ("Ndoshta do të votoja", "Možda bih glasao/ala", "Might vote"),
                ("Me gjasë nuk do të votoja", "Verovatno ne bih glasao/ala", "Probably would not vote"),
                ("Definitivisht nuk do të votoja", "Definitivno ne bih glasao/ala", "Definitely would not vote")
            ]

            all_pairs = core_pairs + voting_pairs + employment_pairs + satisfaction_pairs + awareness_pairs + frequency_pairs + likert_pairs + demographic_pairs + extra_pairs + reason_pairs + income_pairs + demographic_pairs_answers

            for al, sr, en in all_pairs:
                manual_al_to_sr[clean_label(al)] = sr
                manual_sr_to_al[clean_label(sr)] = capitalize_first(al)
                manual_al_to_en[clean_label(al)] = en
                manual_en_to_al[clean_label(en)] = capitalize_first(al)
                manual_sr_to_en[clean_label(sr)] = en
                manual_en_to_sr[clean_label(en)] = sr

            return {
                ("al", "sr"): manual_al_to_sr,
                ("sr", "al"): manual_sr_to_al,
                ("al", "en"): manual_al_to_en,
                ("en", "al"): manual_en_to_al,
                ("sr", "en"): manual_sr_to_en,
                ("en", "sr"): manual_en_to_sr,
            }

        manual_translations = build_translation_dictionaries()

        def fuzzy_lookup(word, dictionary):
            if not word: return ""
            if word in dictionary: return dictionary[word]
            matches = get_close_matches(word, dictionary.keys(), n=1, cutoff=0.9)
            return dictionary[matches[0]] if matches else ""

        unmatched_terms = []

        def apply_manual(text):
            normalized = clean_label(text)
            translation_dict = manual_translations.get((from_lang, to_lang), {})
            translation = fuzzy_lookup(normalized, translation_dict)
            if translation:
                return capitalize_first(translation)
            if normalized:
                unmatched_terms.append(text)
            return ""

        # ── Extract question code from label text (e.g., "S1. Consent" → "s1") ──
        CODE_PATTERN = re.compile(r"^([A-Za-z]+\d+[a-zA-Z]?)[\.\)\:\s]")

        def extract_code(text):
            if pd.isna(text) or not str(text).strip():
                return None
            m = CODE_PATTERN.match(str(text).strip())
            return m.group(1).lower() if m else None

        # ── Build code-based translation maps from Word extraction ──
        question_translations = {}  # code → translated text
        hint_translations = {}      # code → hint text
        option_translations = {}    # (parent_code, option_pos) → translated text

        current_q_code = None
        option_counter = 0
        for _, row in translated_df.iterrows():
            qtext = row.get("Question Text")
            otext = row.get("Option Text")
            hint = row.get("hint", "")
            is_option = pd.notna(row.get("Option ID"))

            if not is_option and pd.notna(qtext) and str(qtext).strip():
                code = extract_code(qtext)
                if code:
                    current_q_code = code
                    option_counter = 0
                    question_translations[code] = str(qtext).strip()
                    if pd.notna(hint) and str(hint).strip():
                        hint_translations[code] = str(hint).strip()
            elif is_option and pd.notna(otext) and str(otext).strip() and current_q_code:
                option_counter += 1
                option_translations[(current_q_code, option_counter)] = str(otext).strip()

        # ── Match survey questions by code ──
        def get_survey_code(row):
            label_text = row.get(from_label, "")
            return extract_code(label_text)

        stats = {"matched": 0, "total": 0}
        unmatched_q = []

        # ── Merge hints ──
        if to_hint_col and "hint" in translated_df.columns:
            def merge_hints(row):
                code = get_survey_code(row)
                if code and code in hint_translations:
                    return hint_translations[code]
                return row.get(to_hint_col, "")
            survey_df[to_hint_col] = survey_df.apply(merge_hints, axis=1)

        # ── Translate survey questions ──
        def translate_question_auto(row):
            row_type = str(row.get("type", "")).strip().lower()
            if row_type.startswith("begin_group") or row_type.startswith("end_group"):
                return row.get(from_label, "")
            if row_type.startswith("note"):
                return row.get(from_label, "")

            code = get_survey_code(row)
            if code:
                stats["total"] += 1
                if code in question_translations:
                    stats["matched"] += 1
                    return capitalize_first(question_translations[code])
                else:
                    unmatched_q.append(f"{code}: {str(row.get(from_label, ''))[:50]}")

            # Fallback to manual dictionary
            return apply_manual(row.get(from_label, ""))

        survey_df[to_label] = survey_df.apply(translate_question_auto, axis=1)

        st.caption(f"Pyetje të gjetura: {stats['matched']}/{stats['total']} (Word ka {len(question_translations)} pyetje)")
        if unmatched_q:
            with st.expander(f"{len(unmatched_q)} pyetje pa përkthim"):
                for item in unmatched_q:
                    st.text(item)

        # ── Translate choices by matching parent question code + option position ──
        # Build: list_name → [all question codes that use it]
        list_code_map = defaultdict(list)
        for _, row in survey_df.iterrows():
            row_type_raw = str(row.get("type", "")).strip()
            row_type = row_type_raw.lower()
            if "select_one" in row_type or "select_multiple" in row_type:
                code = get_survey_code(row)
                if code:
                    parts = row_type_raw.split()
                    list_name = parts[1] if len(parts) > 1 else None
                    if list_name and code not in list_code_map[list_name]:
                        list_code_map[list_name].append(code)

        # For matrix lists, also try the section header code (e.g., b0 for B1-B5)
        # Extract base section + "0" as a candidate (b1→b0, c1→c0, etc.)
        def get_header_codes(codes):
            headers = set()
            for c in codes:
                m = re.match(r"([a-z]+)\d+", c)
                if m:
                    headers.add(f"{m.group(1)}0")
            return list(headers)

        choices_df["_option_pos"] = choices_df.groupby("list_name").cumcount() + 1

        def translate_choice(row):
            list_name = row.get("list_name")
            pos = row.get("_option_pos")
            source_text = row.get(from_label, "")

            codes = list_code_map.get(list_name, [])
            # Also try header codes (b0, c0, etc.) for matrix questions
            all_codes = codes + get_header_codes(codes)

            if all_codes and pd.notna(pos):
                for code in all_codes:
                    key = (code, int(pos))
                    if key in option_translations:
                        return capitalize_first(option_translations[key])

            # Fallback to manual dictionary
            manual = apply_manual(source_text)
            if manual:
                return manual
            # Copy source text for names, numbers, municipalities, etc.
            return source_text if pd.notna(source_text) else ""

        if from_label in choices_df.columns and to_label in choices_df.columns:
            choices_df[to_label] = choices_df.apply(translate_choice, axis=1)
        else:
            st.warning(f"Kolona '{from_label}' ose '{to_label}' nuk ekziston në fletën 'choices'.")

        # ── Clean up temp columns ──
        choices_df.drop(columns=["_option_pos"], inplace=True)

        output_file = "translated_output.xlsx"
        with pd.ExcelWriter(output_file, engine="openpyxl") as writer:
            survey_df.to_excel(writer, sheet_name="survey", index=False)
            choices_df.to_excel(writer, sheet_name="choices", index=False)
            settings_df.to_excel(writer, sheet_name="settings", index=False)

        base_name = os.path.splitext(original_file.name)[0]
        translated_file_name = f"{base_name}_perkthyer.xlsx"

        st.success("Përkthimi u përfundua me sukses!")
        st.download_button("Shkarko Excelin e Përkthyer", data=open(output_file, "rb").read(), file_name=translated_file_name)
